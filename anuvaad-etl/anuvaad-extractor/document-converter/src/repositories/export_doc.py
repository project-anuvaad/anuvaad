import requests
import config
import os
import json
from utilities import DocumentUtilities,FileUtilities,MODULE_CONTEXT
from anuvaad_auditor.loghandler import log_info, log_exception
from zipfile import ZipFile
from reportlab.pdfgen import canvas
from jsonpath_rw import jsonpath, parse
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle   
from reportlab.lib.enums import TA_CENTER,TA_JUSTIFY       
from reportlab.pdfbase import pdfmetrics      
from reportlab.pdfbase.ttfonts import TTFont   
from reportlab.lib.fonts import addMapping
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
import pdftotext
import pandas as pd
from docx import Document
from docx.shared import Twips, Cm
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement,qn


doc_utils = DocumentUtilities()

def zipfile_creation(filepath):
    zip_file = filepath.split('.')[0] + '.zip'
    with ZipFile(zip_file, 'w') as myzip:
        myzip.write(filepath)
    os.remove(filepath)
    return zip_file.split('/')[-1]

class DocumentExporterRepository(object):

    def __init__(self):
        pass

    # getting document json data from fetch-content end point of content-handler.
    def get_data_from_ocr_content_handler(self, record_id, start_page=0, end_page=0):
        
        
        try:
            headers = {"Content-Type": "application/json"}
            request_url = doc_utils.generate_url(config.OCR_CONTENT_HANDLER_HOST, record_id, 0, 0)
            log_info("Intiating request to fetch data from %s"%request_url, MODULE_CONTEXT)
            response = requests.get(request_url, headers = headers)
            response_data = response.content
            log_info("Received data from fetch-content end point of ocr-content-handler", MODULE_CONTEXT)
            dict_str = response_data.decode("UTF-8")
            dict_json_data = json.loads(dict_str)
            return dict_json_data
        except Exception as e:
            log_exception("Can not fetch content in content handler:{}".format(str(e)), MODULE_CONTEXT, e)

    # converting document json data into pandas dataframes.
    def get_page_paragraphs_lines(self, page):
        try:
            page_paragraphs    = []
            page_lines         = []
    
            if 'regions'in page.keys():
                for para_region in page['regions']:
                    if 'class' in para_region.keys() and 'regions' in para_region.keys():
                        if para_region['class'] in ['PARA','HEADER','FOOTER','TABLE']:
                            lines = []
                            for line_region in para_region['regions']:
                                if 'class' in line_region.keys() and 'regions' in line_region.keys():
                                    if line_region['class'] in ['LINE','CELL']:
                                        words = []
                                        for word_region in line_region['regions']:
                                            if 'class' in word_region.keys() and 'text' in word_region.keys():
                                                if word_region['class'] in ['WORD','CELL_TEXT']:
                                                    words.append(word_region['text'])

                                        lines.append(' '.join(words) + '\n')
                                        page_lines.append({'boundingBox': doc_utils.vertices_to_boundingbox(line_region['boundingBox']['vertices']), 
                                                    'text': ' '.join(words)})
                            page_paragraphs.append({'boundingBox': doc_utils.vertices_to_boundingbox(para_region['boundingBox']['vertices']), 
                                               'text': ''.join(lines)})
            return page_paragraphs, page_lines
        except Exception as e:
            log_exception("Page regions formation error", MODULE_CONTEXT, e)


    def create_pdf(self,record, pdf_filepath, font_name, font_size=40, scale_factor=4):

        doc_utils.load_font('arial-unicode-ms', config.FONT_DIR)
        
        w, h                      = doc_utils.get_page_dimensions(record["pages"][0])
        pagesize                  = (w/scale_factor, h/scale_factor)
        c                         = canvas.Canvas(pdf_filepath, pagesize=pagesize)
  
        for page in record["pages"]:

            paragraphs, lines     = self.get_page_paragraphs_lines(page)
                
            for line in lines:
                boundingBox, text = line['boundingBox'], line['text']
                x, y, _, _        = boundingBox
                y                 = h - y
                doc_utils.draw_line_text(c, x/scale_factor, y/scale_factor, text, 1.75, 105, font_name, font_size/scale_factor)
            c.showPage()
        c.save()
        log_info('PDF created: %s' % (pdf_filepath),MODULE_CONTEXT)
        return pdf_filepath


    def create_pdf_to_text(self,pdf_filepath):
        try:
            with open(pdf_filepath, "rb") as f:
                pages = pdftotext.PDF(f)
        except Exception as e:
            log_exception("Exception on generating txt from pdf :{}".format(str(e)),MODULE_CONTEXT,e)
            return False
        
        text_filename = os.path.splitext(os.path.basename(pdf_filepath))[0] + '.txt'
        text_filepath = os.path.join(os.path.dirname(pdf_filepath), text_filename)

        try:
            with open(text_filepath, 'w') as txt:
                txt.write("\n\n".join(pages))
        except Exception as e:
            log_exception("Exception on generating txt from pdf :{}".format(str(e)),MODULE_CONTEXT,e)
            return False
        return text_filepath

    def get_lines_for_txt_write(self, page):
        try:
            lines = []
            if 'regions'in page.keys():
                for para_region in page['regions']:
                    if 'class' in para_region.keys() and 'regions' in para_region.keys():
                        if para_region['class'] in ['PARA','HEADER','FOOTER']:
                            lines.append(f"{self.edit_google_vision_text(para_region['text'])}\n")
                        elif para_region['class'] == 'TABLE':
                            cell_texts = para_region['text'].split('<END_OF_CELL>')
                            for text in cell_texts:
                                lines.append(f"{self.edit_google_vision_text(text)}\n")

            return lines
        except Exception as e:
            log_exception("Exception while formatting data for txt write", MODULE_CONTEXT, e)
    
    def write_to_txt(self,record,file_name):
        try:
            for page in record["pages"]:
                lines     = self.get_lines_for_txt_write(page)
                with open(file_name,'a') as output_file:
                    for line in lines:
                        output_file.write(line)
                
        except Exception as e:
            log_exception("Exception while writing to txt", MODULE_CONTEXT, e)

    def edit_google_vision_text(self,text):
        """Method to process the text
        - removing extra spaces craeted before or after special characters
        """
        s1=text
        try:
            log_info("Correcting google vision text to remove extra spacing",MODULE_CONTEXT)
            i=0
            while(i<len(text)):
                s1=text
                if text[i] in ["/","।",'।' ,':','|',"," ,'०',"]","-",")","}"] and text[i-1]==" ":    
                    text=text[:i-1]+text[i:]
                    if i > 0 :
                        if text[i-1] in ["-","[","{","/","("] and text[i]==" ":
                            text=text[:i]+text[i+1:]
                elif text[i] in ["-","[","{","/","("] and text[i+1]==" ":
                    text=text[:i+1]+text[i+2:]
                i=i+1
        except Exception as e:
            log_exception("Exception while correcting google vision text", MODULE_CONTEXT, e)
            return s1
        return text
    
    def write_to_docx(self, data, o_filename):
        dfs, p_layout = self.convert_pagedata_into_df(data["pages"])
        o_zipfile = self.docx_creation(dfs, p_layout, o_filename)
        return o_zipfile

    def convert_pagedata_into_df(self, pages):
        dfs = []
        page_layout = {}
        df_columns = [
            'text_top',
            'text_left',
            'text_width',
            'text_height',
            'text',
            'font_size',
            'font_family',
            'font_style',
        ]

        for page in pages:
            if type(page) != dict:
                continue
            df = pd.DataFrame(columns=df_columns)
            page_layout.update(
                {
                    'page_width': page['boundingBox']['vertices'][2]['x']/2.77, # required
                    'page_height': page['boundingBox']['vertices'][2]['y']/2.77, # required
                })

            for p_region in page['regions'][1:]:
                if p_region['class'] not in ['PARA', 'FOOTER']:
                    continue

                for l_region in p_region['regions']:
                    df.loc[len(df)] = [
                        l_region['boundingBox']['vertices'][0]['y'],
                        l_region['boundingBox']['vertices'][0]['x'],
                        l_region['boundingBox']['vertices'][2]['x'] -
                        l_region['boundingBox']['vertices'][0]['x'],
                        l_region['boundingBox']['vertices'][2]['y'] -
                        l_region['boundingBox']['vertices'][0]['y'],
                        ' '.join([w_region['text']
                                  for w_region in l_region['regions']]),
                        l_region['font']['size'],
                        l_region['font']['family'],
                        l_region['font']['style'],
                    ]
            # custom tuners related to digitization
            for i in ['text_top', 'text_left', 'text_width', 'text_height']:
                df[i] = df[i].apply(lambda x: x/2.77)
            df['font_size'] = df['font_size'].apply(lambda x: x/4)
            # clean df
            df.sort_values('text_top', axis=0, ascending=True, inplace=True)
            df = df.reset_index()
            df = df.where(pd.notnull(df), None)
            dfs.append(df)

        log_info("dataframes formed", MODULE_CONTEXT)
        return dfs, page_layout

    def docx_creation(self, dataframes, page_layout, output_filename):
        doc_utils = DocumentUtilities()
        document = Document()
        section = document.sections[-1]
        section.orientation = WD_ORIENT.PORTRAIT

        section.page_width = Cm(doc_utils.get_cms(page_layout['page_width']))
        section.page_height = Cm(doc_utils.get_cms(page_layout['page_height']))

        section.left_margin = Cm(0.1)
        section.right_margin = Cm(0.1)
        section.top_margin = Cm(0.1)
        section.bottom_margin = Cm(0.1)
        document._body.clear_content()

        for index, df in enumerate(dataframes):
            for index, row in df.iterrows():
                if row['text'] == None:
                    continue
                # add paragraph
                paragraph = document.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Cm(
                    doc_utils.get_cms(row['text_left']))
                # set space after/before para
                if index == 0:
                    paragraph_format.space_before = Twips(
                        doc_utils.pixel_to_twips(df.iloc[index]['text_top']))
                if index != df.index[-1] and df.iloc[index + 1]['text_top'] != row['text_top']:
                    pixel = df.iloc[index + 1]['text_top'] - \
                        row['text_top'] - (row['font_size']*2.77)
                    if pixel > 0:
                        paragraph_format.space_after = Twips(
                            doc_utils.pixel_to_twips(pixel))
                    else:
                        paragraph_format.space_after = Twips(0)
                else:
                    paragraph_format.space_after = Twips(0)
                # add text with metadata(font,size) to para
                run = paragraph.add_run()
                run.add_text(row['text'])
                if row['font_family'] != None and "Bold" in row['font_style']:
                    run.bold = True
                font = run.font
                font.complex_script = True
                if row['font_size'] != None:
                    # only works en-chars, digits + symbols
                    font.size = Twips(
                        doc_utils.pixel_to_twips(row['font_size']))
                # custom element to enable/set size for non-english characters
                c = OxmlElement('w:szCs')
                c.set(qn('w:val'), font.element.rPr.sz.attrib.values()[0])
                font.element.rPr.append(c)
            run.add_break(WD_BREAK.PAGE)
        document.save(output_filename)
        return output_filename
