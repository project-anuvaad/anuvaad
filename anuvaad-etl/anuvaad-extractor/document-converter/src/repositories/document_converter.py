import requests
import config
import os
import re
import base64
import json
import pandas as pd
from utilities.utils import DocumentUtilities
from docx import Document
from docx.shared import Pt
from docx.shared import Twips, Cm,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Length
from utilities import MODULE_CONTEXT
from anuvaad_auditor.loghandler import log_info, log_exception
from zipfile import ZipFile
import uuid
import xlsxwriter
from jsonpath_rw import jsonpath, parse
from docx.oxml.shared import OxmlElement,qn

def zipfile_creation(filepath):
    arcname = filepath.replace(f"{config.DATA_OUTPUT_DIR}/","")
    zip_file = filepath.split('.')[0] + '.zip'
    with ZipFile(zip_file, 'w') as myzip:
        myzip.write(filepath,arcname)
    os.remove(filepath)
    return zip_file.split('/')[-1]

class DocumentConversion(object):

    def __init__(self, DOWNLOAD_FOLDER):
        self.DOWNLOAD_FOLDER = DOWNLOAD_FOLDER

    # getting document json data from fetch-content end point of content-handler.
    def get_data_from_content_handler(self, record_id, user_id, start_page=0, end_page=0):
        doc_utils = DocumentUtilities()
        try:
            headers = {"x-user-id" : user_id, "Content-Type": "application/json"}
            request_url = doc_utils.url_generation(config.CONTENT_HANDLER_ENDPOINT, record_id, start_page, end_page)
            log_info("Intiating request to fetch data from %s"%request_url, MODULE_CONTEXT)
            response = requests.get(request_url, headers = headers)
            response_data = response.content
            log_info("Received data from fetch-content end point of content handler", MODULE_CONTEXT)
            dict_str = response_data.decode("UTF-8")
            dict_json_data = json.loads(dict_str)
            return dict_json_data
        except Exception as e:
            log_exception("Can not fetch content in content handler: ", MODULE_CONTEXT, e)

    # converting document json data into pandas dataframes.
    def convert_page_data_into_dataframes(self, pages):
        try:
            dfs = []
            page_layout = {}
            df_columns = [
                'text_top',
                'text_left',
                'text_width',
                'text_height',
                'text',
                'font_size',
                'font_family',
                'font_color',
                'base64',
                'line_count'
            ]

            for page in pages:
                if type(page) != dict or "text_blocks" not in page.keys():
                    continue
                df = pd.DataFrame(columns=df_columns)
                page_layout.update(
                    {
                        'page_width': page['page_width'],
                        'page_height': page['page_height']
                    }
                )

                for text in page['text_blocks']:
                    t_content = ' '.join([i['tgt'].strip() for i in text['tokenized_sentences']])
                    if len(t_content) == 0:
                        t_content = text['text']
                    df.loc[len(df)] = [
                        text['text_top'],
                        text['text_left'],
                        text['text_width'],
                        text['text_height'],
                        t_content,
                        text['font_size'],
                        text['font_family'],
                        text['font_color'],
                        None,
                        len(text["children"]),
                    ]
                # clean df
                df.sort_values('text_top', axis=0, ascending=True, inplace=True)
                df = df.reset_index()
                df = df.where(pd.notnull(df), None)
                dfs.append(df)
            log_info("dataframes formed", MODULE_CONTEXT)
            return dfs, page_layout
        except Exception as e:
            log_exception("dataframe formation error", MODULE_CONTEXT, e)

    # using dataframe of document json data to create docx file of target sentences. 
    def document_creation(self, dataframes, page_layout, record_id):
        try:
            doc_utils = DocumentUtilities()
            document              = Document()
            section               = document.sections[-1]
            section.orientation   = WD_ORIENT.PORTRAIT
            
            section.page_width    = Cm(doc_utils.get_cms(page_layout['page_width']))
            section.page_height   = Cm(doc_utils.get_cms(page_layout['page_height']))

            section.left_margin   = Cm(0.1)
            section.right_margin  = Cm(0.1)
            section.top_margin    = Cm(0.1)
            section.bottom_margin = Cm(0.1)
            document._body.clear_content()
            
            for index, df in enumerate(dataframes):
                for index, row in df.iterrows():
                    if row['text'] == None :
                        continue
                    # add paragraph
                    paragraph                      = document.add_paragraph()
                    paragraph_format               = paragraph.paragraph_format
                    paragraph_format.left_indent   = Cm(doc_utils.get_cms(row['text_left']))
                    paragraph_format.right_indent   = Cm(doc_utils.get_cms(row['text_left'] + doc_utils.get_cms(row['text_width']))-1)#Cm(doc_utils.get_cms(page_layout['page_width'] - (row['text_left'] + row['text_width']))) #1440180
                    # set space after/before para
                    if index == 0:
                        paragraph_format.space_before = Twips(doc_utils.pixel_to_twips(df.iloc[index]['text_top']))
                    if index != df.index[-1] and df.iloc[index + 1]['text_top'] != row['text_top']:
                        pixel = df.iloc[index + 1]['text_top'] - row['text_top'] - (row['font_size']*row['line_count'])
                        if pixel>0:
                            paragraph_format.space_after = Twips(doc_utils.pixel_to_twips(pixel))
                        else:
                            paragraph_format.space_after = Twips(0)
                    else:
                        paragraph_format.space_after = Twips(0)
                    # add text with metadata(font,size) to para
                    run                            = paragraph.add_run()
                    run.add_text(row['text'])
                    if row['font_family'] != None and "Bold" in row['font_family']:
                        run.bold                   = True
                    font                           = run.font
                    font.complex_script            = True
                    if row['font_size'] != None:
                        font.size                      = Pt(row['font_size']/2) # only works en-chars, digits + symbols
                    # custom element to enable/set size for non-english characters
                    c = OxmlElement('w:szCs')
                    c.set(qn('w:val'),font.element.rPr.sz.attrib.values()[0])
                    font.element.rPr.append(c)
                run.add_break(WD_BREAK.PAGE)
            out_filename = os.path.splitext(os.path.basename(record_id.split('|')[0]))[0] + str(uuid.uuid4()) + '_translated_docx.docx'
            output_filepath = os.path.join(self.DOWNLOAD_FOLDER , out_filename)
            document.save(output_filepath)
            out_filename_zip = zipfile_creation(output_filepath)
            log_info("docx file formation done!! filename: %s"%out_filename_zip, MODULE_CONTEXT)
            return out_filename_zip
        except Exception as e:
            log_exception("dataframe to doc formation failed", MODULE_CONTEXT, e)

    # get all tokenised object from document json data into one list.
    def get_tokenized_sentences(self, json_data):
        try:
            jsonpath_expr = parse('$..tokenized_sentences[*]')
            matches       = jsonpath_expr.find(json_data)
            tokenized_sentences = []
            for match in matches:
                tokenized_sentences.append(match.value)
            return tokenized_sentences
        except Exception as e:
            log_exception("Getting only tokenised sentence object failed", MODULE_CONTEXT, e)

    # create xlsx file of source sentence in one column and taget sentences in adjacent column
    def generate_xlsx_file(self, record_id, json_data):
        try:
            out_xlsx_filename = os.path.splitext(os.path.basename(record_id.split('|')[0]))[0] + str(uuid.uuid4()) + '_src_tgt_xlsx.xlsx'
            output_filepath_xlsx = os.path.join(self.DOWNLOAD_FOLDER , out_xlsx_filename)
            workbook = xlsxwriter.Workbook(output_filepath_xlsx)
            worksheet = workbook.add_worksheet()
            worksheet.write('A1', 'Source Sentence') 
            worksheet.write('B1', 'Target Sentence')
            row, column = 1, 0
            tokenised_sentences = self.get_tokenized_sentences(json_data)
            if tokenised_sentences != [] or tokenised_sentences != None:
                for tokenised_sentence in tokenised_sentences:
                    if tokenised_sentence != [] or tokenised_sentence != None:
                        worksheet.write(row, column, tokenised_sentence['src']) 
                        worksheet.write(row, column + 1, tokenised_sentence['tgt']) 
                        row += 1
            workbook.close()
            out_xlsx_filename_zip = zipfile_creation(output_filepath_xlsx)
            log_info("xlsx file write completed!! filename: %s"%out_xlsx_filename_zip, MODULE_CONTEXT)
            return out_xlsx_filename_zip
        except Exception as e:
            log_exception("xlsx file formation failed", MODULE_CONTEXT, e)

    # breaking large sentences into page width fit sentences
    def break_large_sentence(self, sentence, max_char_in_line):
        try:
            sub_str_list = [sentence[i:i+max_char_in_line] for i in range(0, len(sentence), max_char_in_line)]
            for idx, sub_str in enumerate(sub_str_list):
                if idx+1 < len(sub_str_list):
                    if not sub_str.endswith(' ') or not sub_str_list[idx+1].startswith(' '): 
                        sub_str_split = sub_str.split(' ')
                        last_word_sub_str = sub_str_split[-1]
                        next_sub_str_split = sub_str_list[idx+1].split(' ')
                        first_word_sub_str = next_sub_str_split[0]
                        if len(last_word_sub_str) < len(first_word_sub_str):
                            next_sub_str_split[0] = ' ' + last_word_sub_str + first_word_sub_str
                            del sub_str_split[-1]
                        else:
                            sub_str_split[-1] = last_word_sub_str + first_word_sub_str + ' '
                            del next_sub_str_split[0]
                        sub_str_list[idx] = ' '.join(sub_str_split)
                        sub_str_list[idx+1] = ' '.join(next_sub_str_split)
            return sub_str_list
        except Exception as e:
            log_exception("sentence breaking failed for txt file writing", MODULE_CONTEXT, e)

    # create txt file for translated sentences 
    def create_translated_txt_file(self, record_id, dataframes, page_layout):
        try:
            out_translated_txt_filename = os.path.splitext(os.path.basename(record_id.split('|')[0]))[0] + str(uuid.uuid4()) + '_translated_txt.txt'
            output_filepath_txt = os.path.join(self.DOWNLOAD_FOLDER , out_translated_txt_filename)
            out_txt_file_write = open(output_filepath_txt, 'w')
            max_chars_in_line = int(page_layout['page_width']/13) if page_layout['page_width'] else 500
            for idx, df in enumerate(dataframes):
                for idx, row in df.iterrows():
                    if df.iloc[idx]['text'] != None and idx+1 < df.shape[0]:
                        extra_spaces = int((df.iloc[idx]['text_left'] - 50)/13) if df.iloc[idx]['text_left'] else 50
                        write_str = re.sub(r'^', ' '*extra_spaces, df.iloc[idx]['text'])
                        if df.iloc[idx]['text_top'] != df.iloc[idx+1]['text_top']:
                            if len(write_str) < max_chars_in_line:
                                out_txt_file_write.write("%s\n"%write_str)
                            else:
                                sub_string_list = self.break_large_sentence(write_str, max_chars_in_line)
                                for item in sub_string_list:
                                    out_txt_file_write.write("%s\n"%item)
                        else:
                            same_line_index = 0
                            same_line_status = bool(df.iloc[idx]['text_top'] == df.iloc[idx+same_line_index+1]['text_top'])
                            while same_line_status:
                                if idx+same_line_index+1 < df.shape[0]:
                                    
                                    try:
                                        onwards_line_space =    int((df.iloc[idx+same_line_index+1]['text_left'] - df.iloc[idx+same_line_index]['text_left'] \
                                                            - df.iloc[idx+same_line_index]['text_width'])/13)
                                    except:
                                        onwards_line_space = 50

                                    if df.iloc[idx+same_line_index+1]['text'] != None:
                                        write_str += ' '*onwards_line_space + df.iloc[idx+same_line_index+1]['text']
                                        df = df.replace({df.iloc[idx+same_line_index+1]['text'] : None})
                                    else:
                                        write_str += ' '*onwards_line_space + ''
                                    same_line_index += 1
                                    if idx+same_line_index+1 < df.shape[0]:
                                        same_line_status = bool(df.iloc[idx+same_line_index]['text_top'] == df.iloc[idx+same_line_index+1]['text_top'])
                                    else:
                                        same_line_status = False
                                else:
                                    same_line_status = False
                            if len(write_str) < max_chars_in_line:
                                out_txt_file_write.write("%s\n"%write_str)
                            else:
                                sub_string_list = self.break_large_sentence(write_str, max_chars_in_line)
                                for item in sub_string_list:
                                    out_txt_file_write.write("%s\n"%item)
                    elif df.iloc[idx]['text'] != None and idx+1 == df.shape[0]:
                        extra_spaces = int((df.iloc[idx]['text_left'] - 50)/13) if df.iloc[idx]['text_left'] else 50
                        write_str = re.sub(r'^', ' '*extra_spaces, df.iloc[idx]['text'])
                        if len(write_str) < max_chars_in_line:
                            out_txt_file_write.write("%s\n"%write_str)
                        else:
                            sub_string_list = self.break_large_sentence(write_str, max_chars_in_line)
                            for item in sub_string_list:
                                out_txt_file_write.write("%s\n"%item)
            out_txt_file_write.close()
            out_txt_zip = zipfile_creation(output_filepath_txt)
            log_info("txt file write completed!! filename: %s"%out_txt_zip, MODULE_CONTEXT)
            return out_txt_zip
        except Exception as e:
            log_exception("txt file formation failed", MODULE_CONTEXT, e)