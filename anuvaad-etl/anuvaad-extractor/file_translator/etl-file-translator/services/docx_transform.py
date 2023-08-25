import copy
import os
import uuid

from anuvaad_auditor import log_info
from lxml import etree

import config
from docx import Document
from docx.text.paragraph import Paragraph
from enums.tags import DocxTag
from errors.errors_exception import FileErrors
from services.service import common_obj



class DocxTransform(object):
    def __init__(self, input_filename, json_data, is_new_flow=True):
        self.json_data = json_data
        log_info(f"Test32: jsondata = {self.json_data}",None)
        self.is_new_flow = is_new_flow
        log_info(f"Test32: newflow = {self.is_new_flow}",None)
        self.file_name_without_ext = os.path.splitext(input_filename)[0]
        log_info(f"Test32: filenamewithoutext = {self.file_name_without_ext}", None)
        self.file_id = self.file_name_without_ext
        log_info(f"Test32: fileid = {self.file_id}",None)

        # variables to track para, run, word count for "PAGE BREAK LOGIC" **IMPORTANT**
        self.para_count = 0  # Total no of Para indivisual and in table combined for page limit
        self.run_count = 0  # Total number of runs indivisual and in table combined for page limit
        self.word_count = 0  # Total number of words for page limit
        self.page_number = 1

        # Variables to track no of para, table, sdt for determine no of Paragraphs/Table/SDT we interacted.
        self.sequence_para_index = -1  # Sequence wise total number of para iterated
        self.sequence_table_index = 0  # Sequence wise total number of table iterated
        self.sequence_sdt_index = 0
        self.sequence_text_box_content_index = 0

        # Json Structure
        self.outer_struct = common_obj.outer_struct
        self.page_struct = common_obj.page_struct
        self.para_struct = common_obj.para_struct
        self.run_struct = common_obj.run_struct

        # docx document object
        self.document = None

        # json object
        self.base_json = copy.deepcopy(self.outer_struct)
        self.page_list = self.base_json['result']

        # Trans Map
        self.trans_map = None

        # Fonts used in a DOCX Document
        self.fonts_used = set()

    # reading content of docx file
    def read_docx_file(self, input_filename):
        log_info("read_docx_file :: started the reading docx file: %s" % input_filename, self.json_data)
        input_docx_filepath = common_obj.input_path(input_filename)

        input_docx = Document(input_docx_filepath)
        self.document = input_docx
        return input_docx
    
    def write_json_file(self, transformed_obj):
        #out_file_name = config.DOCX_FILE_PREFIX + self.file_name_without_ext + '.json'
        out_file_name = config.DOCX1_FILE_PREFIX + self.file_name_without_ext + '.json'
        return common_obj.write_json_file(out_file_name=out_file_name, transformed_obj=transformed_obj)

    def distribute_over_runs(self, iterable_obj, trans_para):
        common_obj.distribute_over_runs(iterable_obj=iterable_obj, trans_para=trans_para)

    def generate_json_for_run(self, run, file_id='', table='', cell='', row='', sdt='', sdtc='', para_idx='',
                              run_idx='', txbxContent='', sub_para_idx=''):
        new_run_template = copy.deepcopy(self.run_struct)
        new_run_template['text'] = run.text
        new_run_template['block_id'] = common_obj.generate_id(file_id=file_id, table=table, cell=cell, row=row, sdt=sdt,
                                                              sdtc=sdtc, para=para_idx, run=run_idx, sub_para=sub_para_idx,
                                                              txbxContent=txbxContent)
        return new_run_template

    '''generate_json_for_para accept para object and using that it tries to create a json structure'''

    def generate_json_for_para(self, para, file_id='', table='', cell='', row='', para_idx='', sub_para_idx='', sdt='',
                               sdtc='', txbxContent=''):
        new_para_template = copy.deepcopy(self.para_struct)
        # runs = common_obj.get_runs(para, para_obj=True)
        new_para_template['text'] = common_obj.get_para_text(para)
        new_para_template['block_id'] = common_obj.generate_id(file_id=file_id, table=table, cell=cell, row=row,
                                                               sdt=sdt, sdtc=sdtc, para=para_idx, sub_para=sub_para_idx,
                                                               txbxContent=txbxContent)
        return new_para_template

    def generate_para_struct_json_for_run(self, run, file_id='', table='', cell='', row='', sdt='', sdtc='',
                                          para_idx='', run_idx=''):
        new_para_template = copy.deepcopy(self.para_struct)
        new_para_template['text'] = run.text
        new_para_template['block_id'] = common_obj.generate_id(file_id=file_id, table=table, cell=cell, row=row,
                                                               sdt=sdt, sdtc=sdtc, para=para_idx, run=run_idx)
        return new_para_template

    def generate_json_for_page(self, page_number):
        new_page_template = copy.deepcopy(self.page_struct)
        new_page_template['page_no'] = page_number
        return new_page_template

    def get_fonts_from_ct_font_obj(self, ct_font_obj):
        try:
            if ct_font_obj.ascii:
                self.fonts_used.add(ct_font_obj.ascii)
            if ct_font_obj.hAnsi:
                self.fonts_used.add(ct_font_obj.hAnsi)
            if ct_font_obj.cs:
                self.fonts_used.add(ct_font_obj.cs)
            if ct_font_obj.eastAsia:
                self.fonts_used.add(ct_font_obj.eastAsia)

        except Exception as e:
            log_info(f"get_fonts_from_ct_font_obj :: GETTING FONTS FAILED"
                     f"ERROR: {str(e)}", self.json_data)

    def get_fonts_used_in_a_file(self, file_name):
        used_font_set = set()
        font_lis = []
        try:
            if file_name == config.DOCX_DOCUMENT_XML:
                font_lis = self.document._element.xpath('.//w:rFonts')
            elif file_name == config.DOCX_STYLE_XML:
                font_lis = self.document.styles.element.xpath('.//w:rFonts')

            for font_obj in font_lis:
                self.get_fonts_from_ct_font_obj(ct_font_obj=font_obj)

        except Exception as e:
            log_info(f"get_fonts_used_in_a_file :: GETTING FONTS FAILED"
                     f"ERROR: {str(e)}", self.json_data)

    def get_fonts_used_in_theme_files(self):
        try:
            parts = self.document.part.package.parts
            for part in parts:
                if '/word/theme/theme' in str(part.partname):
                    theme_file_blob = part.blob
                    # getting the theme xml files (It is in byte format)
                    xml_tree = etree.fromstring(theme_file_blob)
                    for elt in xml_tree.iter():
                        if elt.tag.endswith('}latin'):
                            typeface = elt.attrib.get('typeface')
                            if typeface:
                                self.fonts_used.add(typeface)
                        elif elt.tag.endswith('}ea'):
                            typeface = elt.attrib.get('typeface')
                            if typeface:
                                self.fonts_used.add(typeface)
                        elif elt.tag.endswith('}cs'):
                            typeface = elt.attrib.get('typeface')
                            if typeface:
                                self.fonts_used.add(typeface)

        except Exception as e:
            log_info(f"get_fonts_used_in_theme_files :: GETTING FONTS FAILED"
                     f"ERROR: {str(e)}", self.json_data)

    def get_all_the_fonts_used(self):
        # Fonts mentioned in style.xml, theme.xml and document.xml
        try:
            self.get_fonts_used_in_a_file(file_name=config.DOCX_DOCUMENT_XML)
            self.get_fonts_used_in_a_file(file_name=config.DOCX_STYLE_XML)
            self.get_fonts_used_in_theme_files()
        except Exception as e:
            log_info(f"get_all_the_fonts_used :: GETTING FONTS FAILED"
                     f"ERROR: {str(e)}", self.json_data)

    def get_valid_fonts_for_input_locale(self, in_locale):
        valid_fonts = []
        valid_locale_fonts = config.ALLOWED_FONTS[in_locale]
        common_fonts = config.ALLOWED_FONTS['common']
        valid_fonts.extend(valid_locale_fonts)
        valid_fonts.extend(common_fonts)
        return valid_fonts

    def check_if_valid_fonts_used(self, in_locale):
        self.get_all_the_fonts_used()
        valid_fonts = self.get_valid_fonts_for_input_locale(in_locale)
        valid_fonts_used = True
        invalid_fonts = []
        for font in self.fonts_used:
            if config.ALLOWED_FONTS and font not in valid_fonts:
                valid_fonts_used = False
                invalid_fonts.append(font)

        if valid_fonts_used is False:
            raise FileErrors("FONTS_NOT_SUPPORTED", f"Fonts used in the file is not supported, unsupported fonts: {str(invalid_fonts)}")
        else:
            log_info(f"check_if_valid_fonts_used :: FONTS USED: {str(self.fonts_used)}", self.json_data)

    def add_new_page_on_limit_exceeded(self):
        if common_obj.is_page_size_exceeded(DOCX=True, para_count=self.para_count, run_count=self.run_count, word_count=self.word_count):
            self.para_count, self.run_count, self.word_count = common_obj.reset_page_limit(para_count=self.para_count, run_count=self.run_count,
                                                                                           word_count=self.word_count)
            self.page_number += 1
            log_info("generate_json_structure :: Page limit exceeded generating new page obj, new page index: %s" % self.page_number, self.json_data)
            self.page_list.append(self.generate_json_for_page(self.page_number))

    def get_parent_tag_list(self, para):
        parent_tag_lis = []
        for parent in para.iterancestors():
            parent_tag_lis.append(parent.tag.split('}')[-1])
            # Getting all the parent tag of a Paragraph
        return parent_tag_lis

    def get_relevant_parent_tag(self, parent_tag_list):

        if DocxTag.SDTCONTENT.value in parent_tag_list:
            return DocxTag.SDTCONTENT.value

        if DocxTag.TXBXCONTENT.value in parent_tag_list:
            return DocxTag.TXBXCONTENT.value

        if DocxTag.TABLE.value in parent_tag_list:
            return DocxTag.TABLE.value

        return DocxTag.PARAGRAPH.value

    def iterate_paras(self, document):
        paras = document._element.xpath('.//w:p')
        for pid, p in enumerate(paras):
            parent_tag_lis = self.get_parent_tag_list(para=p)
            parent_tag = self.get_relevant_parent_tag(parent_tag_list=parent_tag_lis)
            para = Paragraph(p, paras)
            yield parent_tag, para

    def add_new_paragraph_to_json(self, para):
        try:
            for idx, subpara in enumerate(common_obj.get_runs(para, para_obj=True)):
                self.para_count += 1
                json_para = self.generate_json_for_para(para=subpara, file_id=self.file_id,
                                                        para_idx=str(self.sequence_para_index),
                                                        sub_para_idx=str(idx))

                words_no = common_obj.word_count(json_para.get('text'))
                self.word_count += words_no

                for idr, run in enumerate(subpara):
                    self.run_count += 1
                    json_run = self.generate_json_for_run(run=run, file_id=self.file_id,
                                                          para_idx=str(self.sequence_para_index),
                                                          run_idx=str(idr), sub_para_idx=str(idx))
                    json_para['children'].append(json_run)

                self.page_list[self.page_number - 1]['text_blocks'].append(json_para)

        except Exception as e:
            log_info(f"add_new_paragraph :: JSON CREATION FAILED FOR PARAGRAPH:{self.sequence_para_index}"
                     f"ERROR:{str(e)}", self.json_data)

    def add_run_as_para_json(self, para):
        try:
            self.para_count += 1

            for runs in common_obj.get_runs(para, para_obj=True):
                for idr, run in enumerate(runs):
                    self.run_count += 1
                    json_run = self.generate_para_struct_json_for_run(run=run, file_id=self.file_id,
                                                                      para_idx=str(self.sequence_para_index),
                                                                      run_idx=str(idr))
                    words_no = common_obj.word_count(json_run.get('text'))
                    self.word_count += words_no
                    json_run['children'].append(copy.deepcopy(json_run))
                    self.page_list[self.page_number - 1]['text_blocks'].append(json_run)

        except Exception as e:
            log_info(f"add_new_toc_to_json :: ADD NEW PARA JSON FAILED:{self.sequence_para_index}"
                     f"ERROR: {str(e)}", self.json_data)

    def generate_json_structure(self, document):
        self.document = document
        self.page_list.append(self.generate_json_for_page(self.page_number))
        # START# NEW LOGIC TO GET ALL THE PARA AT ONCE
        for parent_tag, para in self.iterate_paras(document=self.document):
            self.sequence_para_index += 1

            if parent_tag in [DocxTag.SDTCONTENT.value]:
                self.add_run_as_para_json(para=para)
            else:
                self.add_new_paragraph_to_json(para=para)

            if parent_tag in [DocxTag.PARAGRAPH.value]:
                # We don't want to break the page if it is a continuous tag like Table, SDT, TextBox
                # If page limit is exceeded add a new page_json, increment the page_number by 1
                self.add_new_page_on_limit_exceeded()

        log_info(f'Generated JSON FILE for file: {self.file_name_without_ext}', self.json_data)
        return self.base_json

    def translate_paragraphs(self, para):
        try:
            runs = common_obj.get_runs(para, para_obj=True)
            if self.is_new_flow is True:
                for idx, run_list in enumerate(runs):
                    para_id = common_obj.generate_id(file_id=self.file_id, para=str(self.sequence_para_index),
                                                     sub_para=str(idx))
                    if para_id in self.trans_map:
                        self.distribute_over_runs(run_list, trans_para=self.trans_map[para_id])
                    else:
                        # raise FileErrors("translate_docx_file:", "PARA ID :{} not found in fetch content".format(para_id))
                        log_info("PARA ID :{} not found in fetch content".format(para_id), self.json_data)
            else:
                run_list = common_obj.merge_runs(runs)
                para_id = common_obj.generate_id(file_id=self.file_id, para=str(self.sequence_para_index))

                if para_id in self.trans_map:
                    self.distribute_over_runs(run_list, trans_para=self.trans_map[para_id])
                else:
                    # raise FileErrors("translate_docx_file:", "PARA ID :{} not found in fetch content".format(para_id))
                    log_info("PARA ID :{} not found in fetch content".format(para_id), self.json_data)

        except Exception as e:
            log_info(f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR PARAGRAPH SEQ:{self.sequence_para_index}"
                     f"ERROR: {str(e)}", self.json_data)

    #translate para - old flow below
    # def translate_paragraphs(self, para):
    #     try:
    #         for idx, subpara in enumerate(common_obj.get_runs(para, para_obj=True)):
    #             para_id = common_obj.generate_id(file_id=self.file_id, para=str(self.sequence_para_index),
    #                                              sub_para=str(idx))
    #             if para_id in self.trans_map:
    #                 self.distribute_over_runs(subpara, trans_para=self.trans_map[para_id])
    #             else:
    #                 #raise FileErrors("translate_docx_file:", "PARA ID :{} not found in fetch content".format(para_id))
    #                 log_info("PARA ID :{} not found in fetch content".format(para_id), self.json_data)
    #     except Exception as e:
    #         log_info(f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR PARAGRAPH SEQ:{self.sequence_para_index}"
    #                  f"ERROR: {str(e)}", self.json_data)

    def translate_run_stored_as_para_json(self, para):
        try:
            for runs in common_obj.get_runs(para, para_obj=True):
                for idr, run in enumerate(runs):
                    run_id = common_obj.generate_id(file_id=self.file_id,
                                                    para=str(self.sequence_para_index),
                                                    run=str(idr))
                    run = common_obj.get_runs(run, run_obj=True)
                    if run_id in self.trans_map:
                        self.distribute_over_runs(run, trans_para=self.trans_map[run_id])
                    else:
                        log_info(f"translate_run_stored_as_para_json: RUN ID :{idr} not found in fetch content for SDT",
                                 self.json_data)
        except Exception as e:
            log_info(f"translate_run_stored_as_para_json :: DISTRIBUTE OVER RUN FAILED FOR :{self.sequence_para_index}"
                     f"ERROR: {str(e)}", self.json_data)

    def translate_docx_file(self, document, trans_map):
        self.document = document
        self.trans_map = trans_map
        log_info("translate_docx_file :: Translation Docx Process started.", self.json_data)

        for parent_tag, para in self.iterate_paras(document=self.document):
            self.sequence_para_index += 1

            if parent_tag in [DocxTag.SDTCONTENT.value]:
                self.translate_run_stored_as_para_json(para=para)
            else:
                self.translate_paragraphs(para=para)

        return self.document

    def write_docx_file(self, document):
        file_name = str(uuid.uuid4()) + '.docx'
        file_out_path = common_obj.input_path(file_name)
        document.save(file_out_path)
        return file_name

    def remove_table_of_content(self):
        pass

    # def add_new_paragraph_to_json_old_logic(self):
    #     try:
    #         if len(self.document.paragraphs) <= self.sequence_para_index:
    #             raise FileErrors("DOCX_PARAGRAPH_DATA_GEN_ERROR", "Paragraph Data mismatched in the Docx")
    #
    #         self.para_count += 1
    #
    #         para = self.document.paragraphs[self.sequence_para_index]
    #         json_para = self.generate_json_for_para(para=para, file_id=self.file_id, para_idx=str(self.sequence_para_index))
    #
    #         words_no = common_obj.word_count(json_para.get('text'))
    #         self.word_count += words_no
    #
    #         for idr, run in enumerate(common_obj.get_runs(para, para_obj=True)):
    #             self.run_count += 1
    #             json_run = self.generate_json_for_run(run=run, file_id=self.file_id, para_idx=str(self.sequence_para_index), run_idx=str(idr))
    #             json_para['children'].append(json_run)
    #
    #         self.page_list[self.page_number - 1]['text_blocks'].append(json_para)
    #         self.sequence_para_index += 1
    #         try:
    #             if config.DOCX_TXT_BOX_CONTENT_GEN:
    #                 self.add_new_text_box_to_json(paragraph=para)
    #         except Exception as e:
    #             log_info(f"add_new_paragraph_to_json :: JSON CREATION FAILED FOR DOCX_TXT_BOX_CONTENT_GEN:{self.sequence_para_index}"
    #                      f"ERROR:{str(e)}", None)
    #
    #
    #     except Exception as e:
    #         log_info(f"add_new_paragraph :: JSON CREATION FAILED FOR PARAGRAPH:{self.sequence_para_index}"
    #                  f"ERROR:{str(e)}", None)
    #         self.sequence_para_index += 1
    #
    # def add_new_table_to_json_old_logic(self):
    #     try:
    #         if len(self.document.tables) <= self.sequence_table_index:
    #             raise FileErrors("DOCX_TABLE_DATA_GEN_ERROR", "Table Data mismatched in the Docx")
    #
    #         table = self.document.tables[self.sequence_table_index]
    #         idt = self.sequence_table_index
    #         self.sequence_table_index += 1
    #         try:
    #             for idr, row in enumerate(table.rows):
    #                 try:
    #                     for idc, cell in enumerate(row.cells):
    #                         for idp, para in enumerate(cell.paragraphs):
    #                             try:
    #                                 self.para_count += 1
    #                                 json_para = self.generate_json_for_para(para=para, file_id=self.file_id, table=str(idt), cell=str(idc),
    #                                                                         row=str(idr),
    #                                                                         para_idx=str(idp))
    #
    #                                 words_no = common_obj.word_count(json_para.get('text'))
    #                                 self.word_count += words_no
    #
    #                                 for id_run, run in enumerate(common_obj.get_runs(para, para_obj=True)):
    #                                     self.run_count += 1
    #                                     json_run = self.generate_json_for_run(run=run, file_id=self.file_id,
    #                                                                           table=str(idt),
    #                                                                           cell=str(idc),
    #                                                                           row=str(idr),
    #                                                                           para_idx=str(idp),
    #                                                                           run_idx=str(id_run))
    #                                     json_para['children'].append(json_run)
    #
    #                                 self.page_list[self.page_number - 1]['text_blocks'].append(json_para)
    #
    #                             except Exception as e:
    #                                 log_info(f"add_new_table :: ADDED NEW TABLE JSON FAILED FOR TABLE SEQ:{idt}, row:{idr}, cell:{idc}, para:{idp} "
    #                                          f"ERROR: {str(e)}", None)
    #                 except Exception as e:
    #                     log_info(f"add_new_table :: ADDED NEW TABLE JSON FAILED FOR TABLE SEQ:{idt}, row:{idr} "
    #                              f"ERROR: {str(e)}", None)
    #         except Exception as e:
    #             log_info(f"add_new_table :: ADDED NEW TABLE JSON FAILED FOR TABLE SEQ:{idt}"
    #                      f"ERROR: {str(e)}", None)
    #     except Exception as e:
    #         log_info(f"add_new_table :: JSON CREATION FAILED FOR TABLE:{self.sequence_table_index}"
    #                  f"ERROR:{str(e)}", None)
    #         self.sequence_table_index += 1
    #
    # def add_new_toc_to_json_old_logic(self, child):
    #     # Accept CT_SDT object
    #     try:
    #         for sdtc, sdt_content in enumerate(child.sdtContent_lst):
    #             for parax, para in enumerate(sdt_content.p_lst):
    #                 try:
    #                     para = Paragraph(para, sdtc)
    #
    #                     for idr, run in enumerate(common_obj.get_runs(para, para_obj=True)):
    #                         self.run_count += 1
    #                         json_run = self.generate_para_struct_json_for_run(run=run, file_id=self.file_id,
    #                                                                           sdt=str(self.sequence_sdt_index),
    #                                                                           sdtc=str(sdtc),
    #                                                                           para_idx=str(parax),
    #                                                                           run_idx=str(idr))
    #                         words_no = common_obj.word_count(json_run.get('text'))
    #                         self.word_count += words_no
    #                         json_run['children'].append(copy.deepcopy(json_run))
    #                         self.page_list[self.page_number - 1]['text_blocks'].append(json_run)
    #
    #                 except Exception as e:
    #                     log_info(
    #                         f"add_new_toc :: ADDED NEW PARA JSON TO TOC FAILED FOR TABLE SDT:{self.sequence_sdt_index}, sdtc:{sdtc}, para:{parax}"
    #                         f"ERROR: {str(e)}", None)
    #
    #         self.sequence_sdt_index += 1
    #     except Exception as e:
    #         log_info(f"add_new_toc :: JSON CREATION FAILED FOR TOC:{self.sequence_sdt_index}"
    #                  f"ERROR:{str(e)}", None)
    #         self.sequence_sdt_index += 1
    #
    # def add_new_text_box_to_json(self, paragraph):
    #     try:
    #         txbxContents = paragraph._element.xpath('.//w:txbxContent')
    #
    #         if not txbxContents:  # Checking if there is any txbxContent tag in a para
    #             return
    #
    #         for txbxContent in txbxContents:
    #             for pid, p in enumerate(txbxContent.p_lst):
    #                 try:
    #                     para = Paragraph(p, txbxContent)
    #                     json_para = self.generate_json_for_para(para=para, file_id=self.file_id,
    #                                                             txbxContent=str(self.sequence_text_box_content_index),
    #                                                             para_idx=str(pid))
    #                     words_no = common_obj.word_count(json_para.get('text'))
    #                     self.word_count += words_no
    #
    #                     for idr, run in enumerate(common_obj.get_runs(para, para_obj=True)):
    #                         self.run_count += 1
    #                         json_run = self.generate_json_for_run(run=run, file_id=self.file_id,
    #                                                               txbxContent=str(self.sequence_text_box_content_index),
    #                                                               para_idx=str(pid), run_idx=str(idr))
    #                         json_para['children'].append(json_run)
    #                     self.page_list[self.page_number - 1]['text_blocks'].append(json_para)
    #                 except Exception as e:
    #                     log_info(f"add_new_text_box :: JSON CREATION FAILED FOR TXT BOX CONTENT:{self.sequence_text_box_content_index}, PARA:{pid}"
    #                              f"ERROR:{str(e)}", None)
    #             self.sequence_text_box_content_index += 1
    #     except Exception as e:
    #         log_info(f"add_new_text_box :: JSON CREATION FAILED FOR TXT BOX CONTENT:{self.sequence_text_box_content_index}"
    #                  f"ERROR:{str(e)}", None)
    #
    # def generate_json_structure_old_logic(self, document):
    #     self.document = document
    #     self.page_list.append(self.generate_json_for_page(self.page_number))
    #
    #     # START# NEW LOGIC TO ITERATE FILE SEQUENCIALLY
    #     for ide, child in enumerate(self.document.element.body):
    #         # If page limit is exceeded add a new page_json, increment the page_number by 1 and return the page_list
    #         self.add_new_page_on_limit_exceeded()
    #
    #         if config.DOCX_PARAGRAPH_GEN and isinstance(child, CT_P):
    #             self.add_new_paragraph_to_json()
    #
    #         elif config.DOCX_TOC_GEN and isinstance(child, CT_SDT):
    #             self.add_new_toc_to_json(child=child)
    #
    #         elif config.DOCX_TABLE_DATA_GEN and isinstance(child, CT_Tbl):
    #             self.add_new_table_to_json()
    #
    #     log_info(f'Generated JSON FILE for file: {self.file_name_without_ext}', None)
    #     return self.base_json

    # def translate_paragraphs_old_logic(self):
    #     for idx, para in enumerate(self.document.paragraphs):
    #         try:
    #             runs = common_obj.get_runs(para, para_obj=True)
    #             para_id = common_obj.generate_id(file_id=self.file_id, para=str(idx))
    #             if para_id in self.trans_map:
    #                 self.distribute_over_runs(runs, trans_para=self.trans_map[para_id])
    #             else:
    #                 raise FileErrors("translate_docx_file:", "PARA ID :{} not found in fetch content".format(para_id))
    #         except Exception as e:
    #             log_info(f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR PARAGRAPH SEQ:{idx}"
    #                      f"ERROR: {str(e)}", None)
    #
    # def translate_tables_old_logic(self):
    #     for idt, table in enumerate(self.document.tables):
    #         try:
    #             for idr, row in enumerate(table.rows):
    #                 try:
    #                     for idc, cell in enumerate(row.cells):
    #                         for idp, para in enumerate(cell.paragraphs):
    #                             try:
    #                                 runs = common_obj.get_runs(para, para_obj=True)
    #                                 para_id = common_obj.generate_id(file_id=self.file_id,
    #                                                                  table=str(idt),
    #                                                                  row=str(idr),
    #                                                                  cell=str(idc),
    #                                                                  para=str(idp))
    #                                 if para_id in self.trans_map:
    #                                     self.distribute_over_runs(runs, trans_para=self.trans_map[para_id])
    #                                 else:
    #                                     raise FileErrors("translate_docx_file:", "PARA ID :{} not found in fetch content".format(para_id))
    #                             except Exception as e:
    #                                 log_info(
    #                                     f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR TABLE SEQ:{idt}, row:{idr}, cell:{idc}, para:{idp} "
    #                                     f"ERROR: {str(e)}", None)
    #                 except Exception as e:
    #                     log_info(f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR TABLE SEQ:{idt}, row:{idr} "
    #                              f"ERROR: {str(e)}", None)
    #         except Exception as e:
    #             log_info(f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR TABLE SEQ:{idt}"
    #                      f"ERROR: {str(e)}", None)
    #
    # def translate_toc_old_logic(self):
    #     try:
    #         sequence_sdt_index = 0
    #         for ide, child in enumerate(self.document.element.body):
    #             if isinstance(child, CT_SDT):
    #                 try:
    #                     for sdtc, sdt_content in enumerate(child.sdtContent_lst):
    #                         for parax, para in enumerate(sdt_content.p_lst):
    #                             try:
    #                                 para = Paragraph(para, sdtc)
    #                                 for idr, run in enumerate(common_obj.get_runs(para, para_obj=True)):
    #                                     run_id = common_obj.generate_id(file_id=self.file_id,
    #                                                                     sdt=str(sequence_sdt_index),
    #                                                                     sdtc=str(sdtc),
    #                                                                     para=str(parax),
    #                                                                     run=str(idr))
    #                                     run = common_obj.get_runs(run, run_obj=run)
    #                                     if run_id in self.trans_map:
    #                                         self.distribute_over_runs(run, trans_para=self.trans_map[run_id])
    #                                     else:
    #                                         log_info(f"translate_docx_file: RUN ID :{idr} not found in fetch content for SDT", None)
    #                             except Exception as e:
    #                                 log_info(
    #                                     f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR TABLE SDT:{sequence_sdt_index}, sdtc:{sdtc}, para:{parax}"
    #                                     f"ERROR: {str(e)}", None)
    #                     sequence_sdt_index += 1
    #                 except Exception as e:
    #                     log_info(
    #                         f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR SDT:{sequence_sdt_index}, child seq: {ide}, ERROR: {str(e)}",
    #                         None)
    #
    #     except Exception as e:
    #         log_info(f"translate_docx_file :: DISTRIBUTE OVER RUN FAILED FOR SDT, ERROR: {str(e)}", None)
    #
    # def translate_docx_file_old_logic(self, document, trans_map):  # TODO
    #     self.document = document
    #     self.trans_map = trans_map
    #     log_info("translate_docx_file :: Translation Docx Process started.", None)
    #
    #     if config.DOCX_PARAGRAPH_GEN and config.DOCX_PARAGRAPH_TRANS:
    #         self.translate_paragraphs()
    #
    #     if config.DOCX_TABLE_DATA_GEN and config.DOCX_TABLE_DATA_TRANS:
    #         self.translate_tables()
    #
    #     if config.DOCX_TOC_GEN and config.DOCX_TABLE_OF_CONTENT_TRANS:
    #         self.translate_toc()
    #     return self.document
