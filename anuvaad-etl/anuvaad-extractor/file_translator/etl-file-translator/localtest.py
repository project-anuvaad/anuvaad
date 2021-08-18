from docx import Document
from docx.text.paragraph import Run
from docx.oxml.text.run import CT_R
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.content import CT_MATH, CT_DRAWING, CT_SMARTTAG, CT_HYPERLINK

file_path = 'T:/Documents/doc_math.docx'
file_path2 = 'T:/Documents/doc_smartTag.docx'

document = Document(file_path2)

#print(document)


# for idx, para in enumerate(document.paragraphs):
#     print('Total Paragraphs {} :: Curent: {} \n'.format(str(len(document.paragraphs)), str(idx)))
#     math = para._element.xpath('.//m:oMath')
#     if math:
#         print(math)

# def get_special_para(paragraph_list, p_type = 'normal'):
#     sp_para_list = []
#     try:
#         for idx, para in enumerate(paragraph_list):
#             math = para._element.xpath('.//m:oMath')
#             #drawing = para._element.xpath('.//w:drawing')
#             #smartTag = para._element.xpath('.//w:smartTag')
#             if math:
#                 p_type = 'math'
#                 run_list = [idx,p_type,para.runs]
#                 sp_para_list.append(run_list)
#             # elif drawing:
#                 # p_type = 'drawing'
#                 # run_list = [idx,p_type,para.runs]
#                 # sp_para_list.append(run_list)
#         print("List of all special paragraphs:")
#         print(sp_para_list)
#     except Exception as e:
#         print("Error")

# def get_runs_from_special_para(list,paragraphs):
#     spl_runs_list = []
#     for idx, para in enumerate(paragraphs):
#         if idx in list:
#             spl_runs_list.append(para.runs)
#     print(spl_runs_list)


# print(int(0x0000026817E88880))
# print(int(0x0000026817E888E0))
# print(int(0x0000026817E88940))

#get_special_para(document.paragraphs)

#print(id(document.paragraphs[0].runs[0]))

#get_runs_from_special_para([14,15,18],document.paragraphs)

#print(document.paragraphs[14]._element)

############################################################################################################################


def is_special_para(paragraph):
    try:
        math = paragraph._element.xpath('.//m:oMath')
        drawing = paragraph._element.xpath('.//w:drawing')
        #smartTag = paragraph._element.xpath('.//w:smartTag')
        if math:
            return True
        elif drawing:
            return True
        # elif smartTag:
        #     return True
        else:
            print("This is not a special paragraph")
    except Exception as e:
        print("Error in validating special paragraph")


def get_special_runs(paragraph):
    runs = []
   
    if is_special_para(paragraph) == True:
        for rid, child in enumerate(paragraph._element):
            if isinstance(child, CT_R):
                run_obj = Run(child, paragraph)
                runs.append(run_obj)           
            elif isinstance(child, CT_MATH):
                runs.append(child)
            elif isinstance(child, CT_DRAWING):
                runs.append(child)
                # for did,drun in enumerate(child.r_lst):
                #     if isinstance(drun, CT_R):
                #             run_obj = Run(drun, paragraph)
                #             runs.append(run_obj)
            # elif isinstance(child, CT_SMARTTAG):
            #     runs.append(child)
            #     for smid,smtrun in enumerate(child.r_lst):
            #         if isinstance(smtrun, CT_R):
            #                 run_obj = Run(smtrun, paragraph)
            #                 runs.append(run_obj)
    else:
        return None
    #print(runs)
    return runs


def split_special_runs(paragraph):
    split_runs = []
    temp_list = []
    run_list = get_special_runs(paragraph)

    if type(run_list) == list:
        for child in run_list:
            if isinstance(child, CT_MATH) or isinstance(child, CT_DRAWING):
                split_runs.append(temp_list)
                #split_runs.append([child])
                temp_list = []   
            else:
                if child in temp_list:
                    pass
                else:
                    temp_list.append(child)

        split_runs.append(temp_list)
        #return split_runs
        print(split_runs)

    else:
        return None


#########################################################################################################################


#split_special_runs(document.paragraphs[6])

#get_special_runs(document.paragraphs[14])

# for idx, para in enumerate(document.paragraphs):
#     print('Total Paragraphs {} :: Curent: {} \n'.format(str(len(document.paragraphs)), str(idx+1)))
#     dr = para._element.xpath('.//w:drawing')
#     math = para._element.xpath('.//m:oMath')
#     smttag = para._element.xpath('.//w:smartTag')
#     if smttag:
#         print(smttag)
    # if math:
    #     print(math)
    # if dr:
    #     print(dr) 

# for i in range(len(document.paragraphs)):
#     split_special_runs(document.paragraphs[i])



def get_runs(iterable_obj, para_obj=False, run_obj=False, run_lst=False):
    if len([i for i in [para_obj, run_obj, run_lst] if i]) > 1:
        raise Exception('::Get Runs:: more than one can not be True')
    if len([i for i in [para_obj, run_obj, run_lst] if not i]) == 0:
        raise Exception('::Get Runs:: All can not be False')
    if para_obj:
        #TEST#
        if is_special_para(iterable_obj) == True:
            split_special_runs(iterable_obj)
        #TEST#
        else:   
            runs = []
            for rid, child in enumerate(iterable_obj._element):
                if isinstance(child, CT_R):
                    run_obj = Run(child, iterable_obj)
                    runs.append(run_obj)
                elif isinstance(child, CT_HYPERLINK):
                    for hlrid, hl_run in enumerate(child.r_lst):
                        if isinstance(hl_run, CT_R):
                            run_obj = Run(hl_run, iterable_obj)
                            runs.append(run_obj)
                elif isinstance(child, CT_SMARTTAG):
                    for hlrid, hl_run in enumerate(child.r_lst):
                        if isinstance(hl_run, CT_R):
                            run_obj = Run(hl_run, iterable_obj)
                            runs.append(run_obj)
            print(runs)
    if run_obj:
        print([iterable_obj])
    if run_lst:
        print(iterable_obj)


#test_obj = document.paragraphs[14]  #MATH#
test_obj2 = document.paragraphs[49] #SMARTTAG#
get_runs(test_obj2, para_obj ='True')

#split_special_runs(test_obj)



